"""
Módulo para la generación de documentos Word del sílabo universitario.
Contiene toda la lógica relacionada con la creación del documento .docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
import os  
from datetime import datetime
import json


def formatear_referencia_libro(libro):
    """
    Formatea una referencia de libro desde los campos individuales
    
    Args:
        libro (dict): Diccionario con campos como autores, titulo, año, editorial
        
    Returns:
        str: Referencia formateada
    """
    try:
        # Extraer autores
        autores_str = ""
        if "autores" in libro and isinstance(libro["autores"], dict):
            autores_list = []
            for key in sorted(libro["autores"].keys()):
                autor = libro["autores"][key]
                if isinstance(autor, dict) and "apellido" in autor:
                    inicial = autor.get("inicial", "")
                    apellido = autor.get("apellido", "")
                    if inicial and apellido:
                        autores_list.append(f"{apellido}, {inicial}")
                    elif apellido:
                        autores_list.append(apellido)
            
            if autores_list:
                if len(autores_list) == 1:
                    autores_str = autores_list[0]
                elif len(autores_list) == 2:
                    autores_str = f"{autores_list[0]} y {autores_list[1]}"
                else:
                    autores_str = ", ".join(autores_list[:-1]) + f" y {autores_list[-1]}"
        
        # Extraer otros campos
        titulo = libro.get("titulo", "Título pendiente")
        año = libro.get("año", "Año pendiente")
        editorial = libro.get("editorial", "Editorial pendiente")
        
        # Formatear referencia
        if autores_str:
            return f"{autores_str} ({año}). {titulo}. {editorial}."
        else:
            return f"Autor pendiente ({año}). {titulo}. {editorial}."
    
    except Exception:
        return "Referencia pendiente"


def formatear_referencia_articulo(articulo):
    """
    Formatea una referencia de artículo desde los campos individuales
    
    Args:
        articulo (dict): Diccionario con campos como autores, titulo, año, revista
        
    Returns:
        str: Referencia formateada
    """
    try:
        # Extraer autores (mismo formato que libros)
        autores_str = ""
        if "autores" in articulo and isinstance(articulo["autores"], dict):
            autores_list = []
            for key in sorted(articulo["autores"].keys()):
                autor = articulo["autores"][key]
                if isinstance(autor, dict) and "apellido" in autor:
                    inicial = autor.get("inicial", "")
                    apellido = autor.get("apellido", "")
                    if inicial and apellido:
                        autores_list.append(f"{apellido}, {inicial}")
                    elif apellido:
                        autores_list.append(apellido)
            
            if autores_list:
                if len(autores_list) == 1:
                    autores_str = autores_list[0]
                elif len(autores_list) == 2:
                    autores_str = f"{autores_list[0]} y {autores_list[1]}"
                else:
                    autores_str = ", ".join(autores_list[:-1]) + f" y {autores_list[-1]}"
        
        # Extraer otros campos
        titulo = articulo.get("titulo", "Título pendiente")
        año = articulo.get("año", "Año pendiente")
        revista = articulo.get("revista", "Revista pendiente")
        volumen = articulo.get("volumen", "")
        numero = articulo.get("numero", "")
        paginas = articulo.get("paginas", "")
        
        # Formatear referencia
        referencia = f"{autores_str if autores_str else 'Autor pendiente'} ({año}). {titulo}. {revista}"
        
        if volumen:
            referencia += f", {volumen}"
            if numero:
                referencia += f"({numero})"
        
        if paginas:
            referencia += f", {paginas}"
        
        referencia += "."
        
        return referencia
    
    except Exception:
        return "Referencia pendiente"


def formatear_referencia_web(web):
    """
    Formatea una referencia web desde los campos individuales
    
    Args:
        web (dict): Diccionario con campos como titulo, url, fecha_acceso
        
    Returns:
        str: Referencia formateada
    """
    try:
        titulo = web.get("titulo", "Título pendiente")
        url = web.get("url", "URL pendiente")
        fecha_acceso = web.get("fecha_acceso", "Fecha pendiente")
        autor = web.get("autor", "")
        
        if autor:
            return f"{autor}. {titulo}. Recuperado de {url}. Acceso: {fecha_acceso}."
        else:
            return f"{titulo}. Recuperado de {url}. Acceso: {fecha_acceso}."
    
    except Exception:
        return "Referencia pendiente"


def insertar_campo(paragraph, tipo_campo):
    """
    Inserta un campo de Word en un párrafo (como PAGE o NUMPAGES)
    
    Args:
        paragraph: El párrafo donde insertar el campo
        tipo_campo (str): Tipo de campo ('PAGE' o 'NUMPAGES')
    """
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = tipo_campo
    
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    
    run = paragraph.add_run()
    run._r.append(fldChar_begin)
    run = paragraph.add_run()
    run._r.append(instrText)
    run = paragraph.add_run()
    run._r.append(fldChar_end)


def agregar_bordes(tabla):
    """
    Agrega bordes a una tabla
    
    Args:
        tabla: La tabla a la que agregar bordes
    """
    tbl = tabla._tbl
    tbl_borders = OxmlElement('w:tblBorders')
    
    tipos_borde = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    
    for tipo in tipos_borde:
        borde = OxmlElement(f'w:{tipo}')
        borde.set(qn('w:val'), 'single')
        borde.set(qn('w:sz'), '4')
        borde.set(qn('w:space'), '0')
        borde.set(qn('w:color'), '000000')
        tbl_borders.append(borde)
    
    tbl_pr = tbl.find('.//w:tblPr', tbl.nsmap)
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    
    tbl_pr.append(tbl_borders)


def crear_encabezado_profesional(datos):
    """
    Crea un documento Word con encabezado profesional para el sílabo
    
    Args:
        datos (dict): Diccionario con los datos del sílabo
        
    Returns:
        Document: Objeto Document de python-docx con el encabezado profesional
    """
    doc = Document()
    seccion = doc.sections[0]
    seccion.header_distance = Inches(0.3)
    encabezado = seccion.header

    tabla = encabezado.add_table(rows=4, cols=6, width=Inches(2.0))
    tabla.autofit = False
    tabla.allow_autofit = False

    anchos = [1.0, 1.0, 1.0, 1.0, 1.0, 1.1]  
    for i, ancho in enumerate(anchos):
        for celda in tabla.columns[i].cells:
            celda.width = Inches(ancho)

    celda_logo = tabla.cell(0, 0)
    celda_logo.merge(tabla.cell(3, 0))
    parrafo = celda_logo.paragraphs[0]
    run = parrafo.add_run()
    try:
        run.add_picture("UC.png", width=Inches(0.75))
    except:
        run.text = "LOGO UC"
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    celda_titulo = tabla.cell(0, 1)
    celda_titulo.merge(tabla.cell(0, 5))
    parrafo = celda_titulo.paragraphs[0]
    run = parrafo.add_run("FORMACIÓN ACADÉMICA Y PROFESIONAL")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo.paragraph_format.space_after = Pt(0)
    parrafo.paragraph_format.space_before = Pt(0)

    for i in range(1, 6):
        sombreado = OxmlElement('w:shd')
        sombreado.set(qn('w:fill'), 'B7D6F0')
        tabla.cell(0, i)._tc.get_or_add_tcPr().append(sombreado)

    celda1_col1 = tabla.cell(1, 1)
    celda1_col1.merge(tabla.cell(1, 2))
    celda1_col1.text = "PROCESO NIVEL 0:"
    celda1_col1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    celda1_col2 = tabla.cell(1, 3)
    celda1_col2.merge(tabla.cell(1, 5))
    celda1_col2.text = "ENSEÑANZA – APRENDIZAJE"
    celda1_col2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    celda1_col2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    celda2_col1 = tabla.cell(2, 1)
    celda2_col1.merge(tabla.cell(2, 2))
    p_registro = celda2_col1.paragraphs[0]
    p_registro.clear() 
    run_registro = p_registro.add_run("REGISTRO")
    run_registro.font.size = Pt(8)
    run_registro.font.name = "Arial"
    celda2_col1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_registro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_registro.paragraph_format.left_indent = Cm(-3.2)
    p_registro.paragraph_format.right_indent = Cm(0)
    
    celda2_col2 = tabla.cell(2, 3)
    celda2_col2.merge(tabla.cell(2, 5))
    celda2_col2.text = "FORMATO DE SÍLABO"
    celda2_col2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    celda2_col2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    celda3_col1 = tabla.cell(3, 1)
    celda3_col1.merge(tabla.cell(3, 2))
    p_codigo = celda3_col1.paragraphs[0]
    run_codigo = p_codigo.add_run(f"Código: {datos.get('SLB-COD', '')}")
    run_codigo.font.size = Pt(8)
    run_codigo.font.name = "Arial"
    celda3_col1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p_codigo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_codigo.paragraph_format.space_before = Pt(12)
    p_codigo.paragraph_format.space_after = Pt(6)
    p_codigo.paragraph_format.left_indent = Cm(0)
    p_codigo.paragraph_format.right_indent = Cm(1.5)

    celda3_v = tabla.cell(3, 3)
    p_v = celda3_v.paragraphs[0]
    run_v = p_v.add_run(f"Versión: {datos.get('SLB-VER', '')}")
    run_v.font.size = Pt(9)
    run_v.font.name = "Arial"
    p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_v.paragraph_format.space_before = Pt(6)
    p_v.paragraph_format.space_after = Pt(0)
    celda3_v.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    celda3_f = tabla.cell(3, 4)
    p_f = celda3_f.paragraphs[0]
    run_f = p_f.add_run(f"Fecha: {datos.get('SLB-DATE', '')}")
    run_f.font.size = Pt(9)
    run_f.font.name = "Arial"
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_f.paragraph_format.space_before = Pt(6)
    p_f.paragraph_format.space_after = Pt(0)
    celda3_f.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    celda3_p = tabla.cell(3, 5)
    p_p = celda3_p.paragraphs[0]
    p_p.add_run("Página: ")
    insertar_campo(p_p, 'PAGE')
    p_p.add_run(" de ")
    insertar_campo(p_p, 'NUMPAGES')
    p_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_p.paragraph_format.space_before = Pt(6)
    p_p.paragraph_format.space_after = Pt(0)
    celda3_p.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for fila in tabla.rows[1:4]:
        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                parrafo.paragraph_format.space_before = Pt(6)
                parrafo.paragraph_format.space_after = Pt(0)
                for run in parrafo.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Arial"
                texto = parrafo.text.strip()
                if texto in ["PROCESO NIVEL 0:", "REGISTRO", f"Código: {datos.get('SLB-COD', '')}"]:
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for fila in tabla.rows[1:4]:
        fila.height = Inches(0.2)

    agregar_bordes(tabla)
 #------------------------------------------------------------------------------------------------------

    programa = datos.get('SLB-PROG', '') or datos.get('maestria', '') or ''
    asignatura = datos.get('SLB-ASSIG', '') or datos.get('asignatura', '') or ''
    semestre = datos.get('SLB-SEM', '') or datos.get('semestre', '') or ''
    docente = datos.get('SLB-DOC', '') or datos.get('docente', '') or ''
    
    parrafos = [
        "UNIVERSIDAD NACIONAL DEL CALLAO",
        "ESCUELA DE POSGRADO DE LA UNAC",
        "UNIDAD DE POSGRADO DE LA FACULTAD DE INGENIERÍA MECÁNICA Y DE ENERGÍA",
        "", 
        "SILABO", 
        f"PROGRAMA DE POSGRADO: ",
        f"MAESTRIA EN {programa.upper()}" if programa else "MAESTRIA EN [PROGRAMA]",
        f"ASIGNATURA: {asignatura.upper()}" if asignatura else "ASIGNATURA: [ASIGNATURA]",
        f"SEMESTRE ACADÉMICO: 2025 - {semestre.upper()}" if semestre else "SEMESTRE ACADÉMICO: 2025 - [SEMESTRE]",
        f"DOCENTE: {docente.upper()}" if docente else "DOCENTE: [DOCENTE]",
        "",
        "CALLAO, PERÚ",
        "2025",
    ]

    for idx, texto in enumerate(parrafos):
        if texto == "SILABO":
            p = doc.paragraphs[-1]
            run = p.add_run(texto)
        else:
            p = doc.add_paragraph()
            run = p.add_run(texto)

        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(6)

        if texto == "UNIVERSIDAD NACIONAL DEL CALLAO":
            run.font.size = Pt(22)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif texto in ("CALLAO, PERÚ", "2025"):
            run.font.size = Pt(18)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif texto == "UNIDAD DE POSGRADO DE LA FACULTAD DE INGENIERÍA MECÁNICA Y DE ENERGÍA":
            run.font.size = Pt(14)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)

            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p_img.add_run().add_picture('UC.png', width=Inches(1.8))
            except:
                p_img.add_run("LOGO UNAC")

        elif texto == "SILABO":
            run.font.size = Pt(36)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)

        elif texto.startswith("MAESTRIA EN"):
            run.font.size = Pt(14)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(18)

        elif texto.startswith(("ASIGNATURA", "SEMESTRE ACADÉMICO", "DOCENTE")):
            run.font.size = Pt(14)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        else:
            run.font.size = Pt(14)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#------------------------------------------------------------------------------------------------------v
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    run.add_break()
    run = p.add_run("SILABO")
    run.font.name = 'Times New Roman'
    run.bold = True
    run.font.size = Pt(16)
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run2 = p2.add_run("I. DATOS GENERALES")
    run2.font.name = 'Times New Roman'
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0, 0, 0)
    p2.paragraph_format.left_indent = Inches(0)
    p2.paragraph_format.right_indent = Inches(0)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    p2.style = 'Heading 1'

    horas_teoria = int(datos.get('SLB-HORAT', 0) or datos.get('horas_teoria', 0) or 0)
    horas_practica = int(datos.get('SLB-HORAP', 0) or datos.get('horas_practica', 0) or 0)
    creditos = int(datos.get('CRED', 0) or datos.get('creditos', 0) or 0)

    horas_totales = horas_teoria + horas_practica

    if creditos > 0 and horas_totales != creditos * 4:
        print(f"Advertencia: Las horas totales ({horas_totales}) no coinciden con los créditos asignados ({creditos * 4}).")

    asignatura = datos.get('SLB-ASSIG', '') or datos.get('asignatura', '') or '[ASIGNATURA]'
    codigo_programa = datos.get('SLB-CODPROG', '') or datos.get('codigo_programa', '') or '[CÓDIGO]'
    caracter = datos.get('SLB-CARACTER', '') or datos.get('caracter', '') or '[CARÁCTER]'
    semestre = datos.get('SLB-SEM', '') or datos.get('semestre', '') or '[SEMESTRE]'
    sesiones = datos.get('SESIONES', '') or datos.get('sesiones', '') or '[SESIONES]'
    semanas = datos.get('SEMANAS', '') or datos.get('semanas', '') or '[SEMANAS]'
    docente = datos.get('SLB-DOC', '') or datos.get('docente', '') or '[DOCENTE]'
    correo = datos.get('SLB-CORREO', '') or datos.get('correo', '') or '[CORREO]'
    modalidad = datos.get('SLB-MODALIDAD', '') or datos.get('modalidad', '') or '[MODALIDAD]'

    tabla_datos = [
        ["1.1", "Asignatura", asignatura.title()],
        ["1.2", "Código", codigo_programa.upper()],
        ["1.3", "Carácter", caracter],  
        ["1.4", "Requisito (nombre y código)", "Ninguno"],
        ["1.5", "Ciclo", "I"],
        ["1.6", "Semestre académico", semestre],
        ["1.7", "Número de horas de clase", f"{horas_totales} horas semanales."],
        ["", "Horas de teoría", f"{horas_teoria} horas semanales."],
        ["", "Horas de práctica", f"{horas_practica} horas semanales."],
        ["1.8", "Número de créditos", str(creditos)],
        ["1.9", "Duración", f"{sesiones} sesiones {semanas} semanas"],
        ["1.10", "Docente(s)", docente.title()],
        ["", "Correo electrónico institucional", correo],
        ["1.11", "Modalidad", modalidad]
    ]

    tabla = doc.add_table(rows=len(tabla_datos), cols=3)
    tabla.style = "Table Grid"
    tabla.autofit = False

    widths = (Inches(0.5), Inches(2.0), Inches(3.6))
    for row in tabla.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    for i, fila in enumerate(tabla_datos):
        for j, valor in enumerate(fila):
            cell = tabla.cell(i, j)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT 
            run = paragraph.add_run(str(valor))
            run.font.name = 'Times New Roman' 
            run.font.size = Pt(11)  
            paragraph.paragraph_format.line_spacing = Pt(14)  
            paragraph.paragraph_format.left_indent = Cm(0.1)  

    tabla.cell(6, 0).merge(tabla.cell(8, 0))
    tabla.cell(11, 0).merge(tabla.cell(12, 0))

    def set_table_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        for el in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(el)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    set_table_borders(tabla)

    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run3 = p3.add_run("II. SUMILLA")
    run3.font.name = 'Times New Roman'
    run3.bold = True
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0, 0, 0)
    p3.paragraph_format.left_indent = Inches(0)
    p3.paragraph_format.right_indent = Inches(0)
    p3.paragraph_format.space_before = Pt(12)
    p3.paragraph_format.space_after = Pt(0)
    p3.style = 'Heading 1'

    proposito = datos.get('SLB-PROPOSITO', '') or datos.get('proposito', '') or '[PROPÓSITO DE LA ASIGNATURA]'

    p4 = doc.add_paragraph()
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run4 = p4.add_run(
        f"La asignatura de {asignatura.title()} pertenece al módulo curricular de estudios de especialidad, es de naturaleza teórico-práctico y de de carácter {caracter.lower()}, tiene por propósito {proposito} "
    )
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(0, 0, 0)
    p4.paragraph_format.first_line_indent = Inches(0)
    p4.paragraph_format.space_before = Pt(6)
    p4.paragraph_format.space_after = Pt(0)

    p4_1 = doc.add_paragraph()
    p4_1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run4_1 = p4_1.add_run("La asignatura se organiza en cuatro unidades de aprendizaje:")
    run4_1.font.name = 'Times New Roman'
    run4_1.font.size = Pt(12)
    run4_1.font.color.rgb = RGBColor(0, 0, 0)
    p4_1.paragraph_format.first_line_indent = Inches(0)
    p4_1.paragraph_format.space_before = Pt(0)
    p4_1.paragraph_format.space_after = Pt(0)

    p4_2 = doc.add_paragraph()
    p4_2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_unidad(texto, descripcion):
        r = p4_2.add_run(texto)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 0)

        descripcion = descripcion.strip()
        if descripcion:
            descripcion = descripcion[0].upper() + descripcion[1:]

        r2 = p4_2.add_run(descripcion + "\n")
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0, 0, 0)

    if "temas_por_unidad" in datos:
        for unidad, tema in datos["temas_por_unidad"].items():
            add_unidad(f"{unidad}: ", tema)
    else:
        add_unidad("Unidad I: ", "Tema pendiente")
        add_unidad("Unidad II: ", "Tema pendiente")
        add_unidad("Unidad III: ", "Tema pendiente")
        add_unidad("Unidad IV: ", "Tema pendiente")
#----------------------------------------------------------------------------------------------------------------------------------------
    def capitalizar_despues_de_dos_puntos(texto):
        partes = texto.split(':')
        if len(partes) > 1:
            parte_despues = partes[1].strip()
            if parte_despues:
                partes[1] = ' ' + parte_despues[0].upper() + parte_despues[1:]
        return ':'.join(partes)

    def crear_parrafo(doc, texto, negrita=False, tamano=12, sangria_izq=0, sangria_der=0,
                        espaciado_antes=0, espaciado_despues=0, justificado=False, estilo=None):
            parrafo = doc.add_paragraph()
            parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY if justificado else WD_PARAGRAPH_ALIGNMENT.LEFT
            run = parrafo.add_run(texto)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(tamano)
            run.bold = negrita
            run.font.color.rgb = RGBColor(0, 0, 0)
            parrafo.paragraph_format.left_indent = Inches(sangria_izq)
            parrafo.paragraph_format.right_indent = Inches(sangria_der)
            parrafo.paragraph_format.space_before = Pt(espaciado_antes)
            parrafo.paragraph_format.space_after = Pt(espaciado_despues)
            if estilo:
                parrafo.style = estilo
            return parrafo

    crear_parrafo(doc, "III. COMPETENCIA(S) DEL PERFIL DE EGRESO Y RESULTADO DE APRENDIZAJE",
                negrita=True, tamano=12, estilo='Heading 1')

    crear_parrafo(doc, "RESULTADO DE APRENDIZAJE CAPACIDAD (ES)",
                negrita=True, tamano=12, sangria_izq=0.3, espaciado_antes=6)

    crear_parrafo(doc, "3.1\tResultado de aprendizaje general (Competencias generales)",
                negrita=True, tamano=12, sangria_izq=0.1, espaciado_antes=12, estilo='Heading 2')

    crear_parrafo(doc, "RAG1 (CG1). Trabajo en equipo.",
                negrita=True, tamano=12, sangria_izq=0.5)

    crear_parrafo(doc,
        "Aplica los conceptos básicos de planificación, organización, dirección y control estratégico en la industria, empresas y organizaciones en general, generando su desarrollo y rentabilidad. "
        "Participa en equipo con diversas funciones, inter y multidisciplinarias, para el logro de metas, bajo presión, respetando la diversidad de opiniones con actitud ética y colaborativa.",
        justificado=True, tamano=11, sangria_izq=0.6, espaciado_antes=3, espaciado_despues=12)

    crear_parrafo(doc, "3.2\tResultado de aprendizaje específico (Competencias específicas)",
                negrita=True, tamano=12, sangria_izq=0.1, espaciado_antes=12, estilo='Heading 2')

    competencias = datos.get("competencias_especificas", [])
    for codigo, descripcion in competencias:
        parrafo_codigo_titulo = doc.add_paragraph()
        parrafo_codigo_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  
        parrafo_codigo_titulo.paragraph_format.left_indent = Cm(1.3)  
        parrafo_codigo_titulo.paragraph_format.space_before = Pt(6)   
        parrafo_codigo_titulo.paragraph_format.space_after = Pt(0)    

        descripcion = capitalizar_despues_de_dos_puntos(descripcion.strip())
        run_codigo_titulo = parrafo_codigo_titulo.add_run(f"{codigo} {descripcion.split(':')[0]}:")
        run_codigo_titulo.font.name = 'Times New Roman'
        run_codigo_titulo.font.size = Pt(12)
        run_codigo_titulo.bold = True

        if len(descripcion.split(':')) > 1:
            parrafo_descripcion = doc.add_paragraph()
            parrafo_descripcion.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  
            run_desc = parrafo_descripcion.add_run(descripcion.split(':')[1].strip())
            run_desc.font.name = 'Times New Roman'
            run_desc.font.size = Pt(11)
            run_desc.font.color.rgb = RGBColor(0, 0, 0)
            run_desc.bold = False

            parrafo_descripcion.paragraph_format.left_indent = Inches(0.5)
            parrafo_descripcion.paragraph_format.right_indent = Inches(0)
            parrafo_descripcion.paragraph_format.space_before = Pt(6)
            parrafo_descripcion.paragraph_format.space_after = Pt(0)

    crear_parrafo(doc, "3.3\tProducto(s) o actividad(es) de aprendizaje evaluados",
                negrita=True, tamano=12, sangria_izq=0.1, espaciado_antes=12, estilo='Heading 2')

    actividades = datos.get("productos_actividades")
    if not actividades or not isinstance(actividades, list) or not actividades:
        actividades = [
            ("PA1 (C1).", "Producto pendiente"),
            ("PA2 (C2).", "Producto pendiente"),
            ("PA3 (C3).", "Producto pendiente"),
            ("PA4 (C4).", "Producto pendiente")
        ]

    for codigo, texto in actividades:
        parrafo_codigo_titulo = doc.add_paragraph()
        parrafo_codigo_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  
        parrafo_codigo_titulo.paragraph_format.left_indent = Cm(1.3)  
        parrafo_codigo_titulo.paragraph_format.right_indent = Cm(0)   
        parrafo_codigo_titulo.paragraph_format.space_before = Pt(6)  
        parrafo_codigo_titulo.paragraph_format.space_after = Pt(0)   

        texto = capitalizar_despues_de_dos_puntos(texto.strip())
        run_codigo_titulo = parrafo_codigo_titulo.add_run(f"{codigo} {texto.split(':')[0]}:")
        run_codigo_titulo.font.name = 'Times New Roman'
        run_codigo_titulo.font.size = Pt(12)
        run_codigo_titulo.bold = True
        run_codigo_titulo.font.color.rgb = RGBColor(0, 0, 0)

        parrafo_descripcion = doc.add_paragraph()
        parrafo_descripcion.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  
        parrafo_descripcion.paragraph_format.space_before = Pt(0)      
        parrafo_descripcion.paragraph_format.space_after = Pt(0)      

        if len(texto.split(':')) > 1:
            run_desc = parrafo_descripcion.add_run(texto.split(':')[1].strip())
            run_desc.font.name = 'Times New Roman'
            run_desc.font.size = Pt(11)
            run_desc.font.color.rgb = RGBColor(0, 0, 0)
            run_desc.bold = False

        parrafo_descripcion.paragraph_format.left_indent = Inches(0.5)
        parrafo_descripcion.paragraph_format.right_indent = Inches(0)
        parrafo_descripcion.paragraph_format.space_before = Pt(6)
        parrafo_descripcion.paragraph_format.space_after = Pt(0)

    crear_parrafo(doc, "IV. METODOLOGÍA",
                negrita=True, tamano=12, estilo='Heading 1', espaciado_antes=6)

    metodologia_texto = (
        "La Universidad Nacional del Callao, Licenciada por la SUNEDU tiene como fin supremo la formación integral del estudiante, quien es el eje central del proceso educativo de formación profesional; "
        "es así como el Modelo Educativo de la UNAC implementa las teorías educativas constructivista y conectivista, y las articula con los componentes transversales del proceso de enseñanza – aprendizaje, "
        "orientando las competencias genéricas y específicas. Este modelo tiene como propósito fundamental la formación holística de los estudiantes y concibe el proceso educativo en la acción y para la acción. "
        "Además, promueve el aprendizaje significativo en el marco de la construcción o reconstrucción cooperativa del conocimiento y toma en cuenta los saberes previos de los participantes con la finalidad que los "
        "estudiantes fortalezcan sus conocimientos y formas de aprendizaje y prosperen en la era digital, en un entorno cambiante de permanente innovación, acorde con las nuevas herramientas y tecnologías de información y comunicación."
    )

    crear_parrafo(doc, metodologia_texto,
                justificado=True, tamano=12, sangria_izq=0.3, espaciado_antes=6, espaciado_despues=6)

    metodologia_texto_2 = (
        "La Facultad de Ingeniería Mecánica y de Energía de la UNAC, en cumplimiento con lo dispuesto en la Resolución Viceministerial Nº085-2020-MINEDU del 01 de abril de 2020, "
        "de manera excepcional y mientras duren las medidas adoptadas por el Gobierno con relación al estado de emergencia sanitario, se impartirá educación remota no presencial haciendo uso de una plataforma virtual educativa: "
        "espacio en donde se imparte el servicio educativo de los cursos, basados en tecnologías de la información y comunicación (TICs)."
    )

    crear_parrafo(doc, metodologia_texto_2,
                justificado=True, tamano=12, sangria_izq=0.3, espaciado_antes=0, espaciado_despues=6)

    metodologia_texto_3 = (
        "La plataforma de la UNAC es el Sistema de Gestión Académico (SGA-UNAC) basado en Moodle, en donde los estudiantes, tendrán a su disposición información detallada del curso: el sílabo, la matriz formativa, "
        "ruta del aprendizaje, guía de entregables calificados, y los contenidos de la clase estructurados para cada sesión educativa. El SGA será complementado con las diferentes soluciones que brinda Google Suite for Education y otras herramientas tecnológicas multiplataforma."
    )

    crear_parrafo(doc, metodologia_texto_3,
                justificado=True, tamano=12, sangria_izq=0.3, espaciado_antes=0, espaciado_despues=6)

    crear_parrafo(doc,
                "Las estrategias metodológicas didáctica para el desarrollo de las sesiones teóricas permiten dos modalidades de aprendizaje en los estudiantes:",
                justificado=True, tamano=12, sangria_izq=0.3, espaciado_antes=0, espaciado_despues=6)

    crear_parrafo(doc, "4.1 Herramientas metodológicas de comunicación síncrona (videoconferencia)", negrita=True, tamano=12, espaciado_antes=6, estilo='Heading 2')

    crear_parrafo(doc,
        "La modalidad síncrona es una forma de aprendizaje basado en el uso de herramientas que permiten la comunicación no presencial y en tiempo real entre el docente y los estudiantes.",
        justificado=True, tamano=12, sangria_izq=0.3, espaciado_antes=6)

    crear_parrafo(doc,
        "Dentro de la modalidad síncrona, se hará uso de:",
        justificado=True, tamano=12, sangria_izq=0.3)

    crear_parrafo(doc,
        "Clases dinámicas e interactivas (virtuales): el docente genera permanentemente expectativa por el tema a través de actividades que permiten vincular los saberes previos con el nuevo conocimiento, promoviendo la interacción mediante el diálogo y debate sobre los contenidos.",
        justificado=True, tamano=12, sangria_izq=0.3, espaciado_antes=6)

    crear_parrafo(doc,
        "Talleres de aplicación (virtuales): el docente genera situaciones de aprendizaje para la transferencia de los aprendizajes a contextos reales o cercanos a los participantes que serán retroalimentados en clase.",
        justificado=True, tamano=12, sangria_izq=0.3, espaciado_antes=6)

    crear_parrafo(doc,
        "Tutorías (virtuales): Para facilitar la demostración, presentación y corrección de los avances del informe final de investigación.",
        justificado=True, tamano=12, sangria_izq=0.3, espaciado_antes=6)

    crear_parrafo(doc, "4.2 Herramientas metodológicas de modalidad asíncrona", negrita=True, tamano=12, espaciado_antes=6, estilo='Heading 2')

    crear_parrafo(doc,
        "Forma de aprendizaje basado en el uso de herramientas que posibilitan el intercambio de mensajes e información entre los estudiantes y el docente en tiempo diferido y sin interacción instantánea.",
        justificado=True, tamano=12, sangria_izq=0.3, espaciado_antes=6)

    crear_parrafo(doc,
        "Dentro de la modalidad asincrónica se hará uso de metodologías colaborativas tales como:",
        justificado=True, tamano=12, sangria_izq=0.3)

    crear_parrafo(doc,
        "- Aprendizaje Orientado a Proyectos - AOP (virtual): Permite que el estudiante adquiriera conocimientos y competencias mediante la ejecución de su proyecto de investigación, para dar respuesta a problemas del contexto.",
        justificado=True, tamano=12, sangria_izq=0.3)

    crear_parrafo(doc,
        "- Portafolio de Evidencias Digital: Permite dar seguimiento a la organización y presentación de evidencias de investigación y recopilación de información para poder observar, contrastar, sugerir, incentivar, preguntar.",
        justificado=True, tamano=12, sangria_izq=0.3)

    crear_parrafo(doc,
        "- Foro de investigación: se realizarán foros de debate, a partir de un reactivo sobre el tema de la sesión de aprendizaje.",
        justificado=True, tamano=12, sangria_izq=0.3)

    crear_parrafo(doc,
        "- Aprendizaje Basado en Problemas (ABP)",
        justificado=True, tamano=12, sangria_izq=0.3)

    crear_parrafo(doc,
        "- Aula invertida",
        justificado=True, tamano=12, sangria_izq=0.3)

    crear_parrafo(doc,
        "- Retroalimentación",
        justificado=True, tamano=12, sangria_izq=0.3)

    def crear_parrafo_estandarizado(doc, texto, negrita=False, tamano=12, espaciado_antes=0, espaciado_despues=0, estilo=None, sangria_izq=0.3, justificado=False):
        crear_parrafo(doc, texto, negrita=negrita, tamano=tamano, espaciado_antes=espaciado_antes, espaciado_despues=espaciado_despues, estilo=estilo, sangria_izq=sangria_izq, justificado=justificado)

    crear_parrafo_estandarizado(doc, "EJES TRANSVERSALES", negrita=True, tamano=12, espaciado_antes=6)

    crear_parrafo_estandarizado(doc, "4.2.1 INVESTIGACIÓN.", negrita=True, tamano=12, espaciado_antes=6, estilo='Heading 2')

    investigacion_contenido = [
        "- Fichas de investigación",
        "- Búsqueda bibliográfica",
        "- Análisis y síntesis de información",
        "- Organización y sistematización de la información",
        "- Uso de referencias",
        "- Estilos de redacción"
    ]

    for item in investigacion_contenido:
        crear_parrafo_estandarizado(doc, item, justificado=True, tamano=12)

    crear_parrafo_estandarizado(doc, "RESPONSABILIDAD SOCIAL", negrita=True, tamano=12, espaciado_antes=6)

    responsabilidad_social_texto = (
        "Proyectos alineados a la problemática social mediante los entregables de acuerdo a cada asignatura según el programa académico y de acuerdo al proyecto de responsabilidad social de la Facultad."
    )
    crear_parrafo_estandarizado(doc, responsabilidad_social_texto, justificado=True, tamano=12)
#------------------------------------------------------------------------------------------------------
    def obtener_fechas_unidades(cronograma_data):
        fechas = {}
        if not cronograma_data or not isinstance(cronograma_data, dict):
            return fechas
        
        # Manejar la estructura del cronograma del usuario
        cronograma_array = cronograma_data.get("cronograma", [])
        if not cronograma_array:
            return fechas
            
        # Agrupar sesiones por unidad
        unidades_sesiones = {}
        for sesion in cronograma_array:
            if isinstance(sesion, dict) and sesion.get("unidad") and sesion.get("fecha"):
                unidad_num = sesion.get("unidad")
                unidad_key = f"Unidad {['I', 'II', 'III', 'IV'][unidad_num - 1]}"
                
                if unidad_key not in unidades_sesiones:
                    unidades_sesiones[unidad_key] = []
                
                # Convertir fecha de formato DD/MM/YYYY a YYYY-MM-DD para ordenar
                fecha_str = sesion.get("fecha", "")
                if fecha_str and fecha_str != "Fecha pendiente":
                    unidades_sesiones[unidad_key].append(fecha_str)
        
        # Obtener fechas de inicio y término para cada unidad
        for unidad, fechas_sesiones in unidades_sesiones.items():
            if fechas_sesiones:
                # Ordenar las fechas (asumiendo formato DD/MM/YYYY)
                fechas_ordenadas = sorted(fechas_sesiones)
                fechas[unidad] = {
                    "inicio": fechas_ordenadas[0],
                    "termino": fechas_ordenadas[-1]
                }
            else:
                fechas[unidad] = {"inicio": "Fecha pendiente", "termino": "Fecha pendiente"}
                
        return fechas

    def obtener_temarios_sesiones(datos_completos):
        """
        Obtiene los temarios de las sesiones desde los datos de sesiones
        """
        temarios = {}
        if not datos_completos:
            return temarios
            
        sesiones_data = datos_completos.get("sesiones", {})
        if not isinstance(sesiones_data, dict):
            return temarios
            
        unidades_sesiones = sesiones_data.get("unidades_sesiones", [])
        if not unidades_sesiones:
            return temarios
            
        # Procesar cada unidad de sesiones
        for unidad_sesiones in unidades_sesiones:
            if not isinstance(unidad_sesiones, dict):
                continue
                
            unidad_numero = unidad_sesiones.get("unidad_numero", 0)
            sesiones = unidad_sesiones.get("sesiones", [])
            
            for sesion in sesiones:
                if isinstance(sesion, dict):
                    numero_sesion = sesion.get("numero_sesion", 0)
                    temario = sesion.get("temario", "Tema pendiente")
                    temarios[numero_sesion] = temario
                    
        return temarios

    doc.add_page_break()

    crear_parrafo_estandarizado(doc, "\nV. ORGANIZACIÓN DE LAS UNIDADES DE APRENDIZAJE", negrita=True, tamano=12, espaciado_antes=18, espaciado_despues=6, estilo='Heading 1')

    fechas_unidades = obtener_fechas_unidades(datos.get("cronograma", {}))
    temas_por_unidad = datos.get("temas_por_unidad", {})
    cronograma_data = datos.get("cronograma", {})
    cronograma = cronograma_data.get("cronograma", [])
    
    # Obtener temarios de sesiones
    temarios_sesiones = obtener_temarios_sesiones(datos)
    
    # Convertir el cronograma del usuario a la estructura esperada
    cronograma_estructurado = {}
    if cronograma:
        for sesion in cronograma:
            if isinstance(sesion, dict) and sesion.get("unidad"):
                unidad_num = sesion.get("unidad")
                unidad_key = f"Unidad {['I', 'II', 'III', 'IV'][unidad_num - 1]}"
                
                if unidad_key not in cronograma_estructurado:
                    cronograma_estructurado[unidad_key] = []
                
                cronograma_estructurado[unidad_key].append({
                    "id": sesion.get("sesion", 1),
                    "fecha": sesion.get("fecha", "Fecha pendiente"),
                    "horas": str(sesion.get("horas", "Hora pendiente"))
                })
        cronograma = cronograma_estructurado
    
    # Asegurar que competencias tenga el formato correcto
    competencias = datos.get("competencias_especificas", [])
    if not competencias or not isinstance(competencias, list):
        competencias = [
            ("RAE1 (CE1).", "Competencia pendiente"),
            ("RAE2 (CE2).", "Competencia pendiente"),
            ("RAE3 (CE3).", "Competencia pendiente"),
            ("RAE4 (CE4).", "Competencia pendiente")
        ]
    indicadores_por_unidad = datos.get("indicadores_por_unidad", {})
    instrumentos_por_unidad = datos.get("instrumentos_por_unidad", {})
    contador_sesion = 1

    if not cronograma:
        cronograma = {
            "Unidad I": [
                {"id": 1, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 2, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 3, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 4, "fecha": "Fecha pendiente", "horas": "Hora pendiente"}
            ],
            "Unidad II": [
                {"id": 5, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 6, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 7, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 8, "fecha": "Fecha pendiente", "horas": "Hora pendiente"}
            ],
            "Unidad III": [
                {"id": 9, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 10, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 11, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 12, "fecha": "Fecha pendiente", "horas": "Hora pendiente"}
            ],
            "Unidad IV": [
                {"id": 13, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 14, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 15, "fecha": "Fecha pendiente", "horas": "Hora pendiente"},
                {"id": 16, "fecha": "Fecha pendiente", "horas": "Hora pendiente"}
            ]
        }

    for i, (unidad, sesiones) in enumerate(cronograma.items(), start=1):
        def convertir_a_romano(numero):
            romanos = {1: 'I', 2: 'II', 3: 'III', 4: 'IV'}
            return romanos.get(numero, str(numero))

        unidad_romana = f"Unidad {convertir_a_romano(i)}"
        tema_unidad = temas_por_unidad.get(unidad_romana, "Tema no definido").capitalize()
        # Usar la clave correcta para obtener las fechas
        fechas = fechas_unidades.get(unidad, fechas_unidades.get(unidad_romana, {"inicio": "Fecha pendiente", "termino": "Fecha pendiente"}))
        tabla = doc.add_table(rows=5, cols=5)
        tabla.style = 'Table Grid'

        tabla.cell(0, 0).merge(tabla.cell(0, 4)).paragraphs[0].add_run(f"UNIDAD DE APRENDIZAJE N° {i}: {tema_unidad}").bold = True
        tabla.cell(1, 0).merge(tabla.cell(1, 4)).paragraphs[0].add_run(f"Fecha de inicio: {fechas['inicio']}     Fecha de término: {fechas['termino']}").bold = True

        celda3 = tabla.cell(2, 0).merge(tabla.cell(2, 4))
        celda3.paragraphs[0].add_run(f"Resultado de aprendizaje específico {i}:").bold = True

        if i <= len(competencias):
            codigo, descripcion = competencias[i - 1]
            partes_descripcion = descripcion.split(":")
            partes_descripcion = [parte.strip().capitalize() for parte in partes_descripcion]
            p_cod = celda3.add_paragraph()
            p_cod.paragraph_format.left_indent = Inches(0.5)
            p_cod.add_run(f"{codigo} {partes_descripcion[0]}:").bold = True
            if len(partes_descripcion) > 1:
                p_desc = celda3.add_paragraph()
                p_desc.paragraph_format.left_indent = Inches(0.5)
                p_desc.add_run(partes_descripcion[1]).bold = False
        else:
            celda3.add_paragraph("• Competencia no definida para esta unidad.")

        celda4 = tabla.cell(3, 0).merge(tabla.cell(3, 4))
        celda4.paragraphs[0].add_run("Producto de aprendizaje de la unidad:").bold = True

        productos_limpios = []
        for p in datos.get("productos_actividades", []):
            if len(p) == 2 and ":" in p[1]:
                partes = p[1].split(":", 1)
                productos_limpios.append((p[0], partes[0].strip(), partes[1].strip()))
            elif len(p) == 3:
                productos_limpios.append(p)
            else:
                productos_limpios.append((p[0], "Título no definido", "Contenido no definido"))
        datos["productos_actividades"] = productos_limpios

        # Buscar producto con diferentes variaciones del código
        codigo_buscado_1 = f"PA{i}(C{i})"
        codigo_buscado_2 = f"PA{i} (C{i})"
        producto = None
        
        for p in productos_limpios:
            if p[0] == codigo_buscado_1 or p[0] == codigo_buscado_2:
                producto = p
                break
        
        if not producto:
            producto = (f"PA{i} (C{i})", "Título no definido", "Contenido no definido")
        parrafo_prod = celda4.add_paragraph()
        parrafo_prod.paragraph_format.left_indent = Inches(0.5)
        parrafo_prod.add_run(f"{producto[0]} {producto[1].capitalize()}:").bold = True
        if producto[2] != "Contenido no definido":
            celda4.add_paragraph(producto[2].capitalize()).paragraph_format.left_indent = Inches(0.5)

        fila_enc = tabla.rows[4].cells
        fila_enc[0].paragraphs[0].add_run("No. Sesión / Horas Lectivas").bold = True
        fila_enc[1].paragraphs[0].add_run("Tema / actividad").bold = True
        fila_enc[2].paragraphs[0].add_run("Indicador (es) de logro").bold = True
        celda_instrum_header = fila_enc[3].merge(fila_enc[4])
        p_instrum = celda_instrum_header.paragraphs[0]
        p_instrum.add_run("Instrumentos de evaluación").bold = True
        p_instrum.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fila_encabezados = tabla.rows[4]
        for cell in fila_encabezados.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  

        indicador_unidad = indicadores_por_unidad.get(unidad_romana, f"Desarrolla competencias específicas relacionadas con {tema_unidad.lower()}")

        for sesion in sesiones:
            fila = tabla.add_row().cells
            
            # Asegurar que sesion sea un diccionario
            if isinstance(sesion, dict):
                fecha = sesion.get("fecha", "Fecha pendiente")
                horas_sesion = sesion.get("horas", "4")
                sesion_id = sesion.get('id', contador_sesion)
            else:
                # Si sesion no es diccionario, usar valores por defecto
                fecha = "Fecha pendiente"
                horas_sesion = "4"
                sesion_id = contador_sesion
                
            p_sesion = fila[0].paragraphs[0]
            p_sesion.add_run(f"SESIÓN {contador_sesion}").bold = True
            p_sesion.add_run(f"\n{horas_sesion} horas\n{fecha}")
            p_sesion.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sesion.paragraph_format.space_before = Pt(24)
            
            # Obtener el tema de la sesión desde los datos de sesiones
            tema_sesion = temarios_sesiones.get(contador_sesion, "Tema pendiente")
            
            p_tema = fila[1].paragraphs[0]
            p_tema.clear()
            for linea in tema_sesion.splitlines():
                linea = linea.strip()
                if linea:
                    p_tema.add_run(f"• {linea[0].upper() + linea[1:]}\n")

            fila[3].merge(fila[4])
            contador_sesion += 1
        for row in tabla.rows[5:]: 
            cell = row.cells[0] 
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  

        num_filas = len(tabla.rows)
        if num_filas > 5:
            celda_indicador = tabla.cell(5, 2).merge(tabla.cell(num_filas - 1, 2))
            celda_indicador.text = indicador_unidad
            celda_indicador.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            celda_indicador.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        celda_instrum = tabla.cell(5, 3).merge(tabla.cell(num_filas - 1, 3))
        p_eval = celda_instrum.paragraphs[0]
        p_eval.clear()
        for idx, instrumento in enumerate(instrumentos_por_unidad.get(unidad_romana, ["Lista de cotejo", "Rúbrica de evaluación"])):
            if idx > 0:
                p_eval = celda_instrum.add_paragraph()
            p_eval.add_run(f"• {instrumento}")
            p_eval.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_eval.paragraph_format.space_before = Pt(0)
            p_eval.paragraph_format.space_after = Pt(0)
        celda_instrum.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for row in tabla.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman' 
                        run.font.size = Pt(11) 

        tabla.autofit = False 
        tabla.allow_autofit = False 

        column_widths = [Inches(1.0), Inches(2.5), Inches(1.5), Inches(1.5), Inches(1.5)]  
        for i, width in enumerate(column_widths):
            for cell in tabla.columns[i].cells:
                cell.width = width

        doc.add_paragraph("")

    crear_parrafo_estandarizado(doc, "(*) Observación: En los programas de Doctorado con asignaturas de duración de módulo de 6 semanas u 8 semanas deberán considerar el número de unidades de aprendizaje según corresponda.", tamano=12, espaciado_antes=6, sangria_izq=0)

    crear_parrafo_estandarizado(doc, "VI. SISTEMA DE EVALUACIÓN DE LA ASIGNATURA", negrita=True, tamano=12, espaciado_antes=18, sangria_izq=0, estilo='Heading 1')
    crear_parrafo_estandarizado(doc, "Se contemplan dentro del sistema de evaluación, evaluación diagnóstica y evaluación sumativa.", tamano=12, espaciado_antes=6)

    crear_parrafo_estandarizado(doc, "REQUISITOS PARA APROBAR LA ASIGNATURA", negrita=True, tamano=12, espaciado_antes=18, estilo='Heading 2')
    crear_parrafo_estandarizado(doc, "Conforme a las disposiciones de los reglamentos de estudios de la Escuela de Posgrado de la Universidad Nacional del Callao, se tendrá a consideración lo siguiente:", tamano=12, espaciado_antes=6)

    requisitos = [
        "● Participación activa en todas las tareas de aprendizaje.",
        "● Asistencia 70% mínimo.",
        "● La escala de calificación es de 0 a 20.",
        "● El estudiante aprueba si su nota promocional es mayor o igual a 14."
    ]

    for requisito in requisitos:
        crear_parrafo_estandarizado(doc, requisito, tamano=12, espaciado_antes=6, sangria_izq=0.4)

    crear_parrafo_estandarizado(doc, "La evaluación del aprendizaje se adecua a la modalidad no presencial, considerando las capacidades y los productos de aprendizaje evaluados descritos para cada unidad. Se evalúa antes, durante y al finalizar el proceso de enseñanza-aprendizaje, considerando la aplicación de los instrumentos de evaluación pertinentes.", tamano=12, espaciado_antes=12, justificado=True)

    crear_parrafo_estandarizado(doc, "Evaluación diagnóstica: Se aplica en la primera sesión de aprendizaje para evaluar los saberes previos necesarios para el desarrollo de la asignatura y la toma de decisiones por parte del docente.", tamano=12, espaciado_antes=6, justificado=True)
#------------------------------------------------------------------------------------------------------

    def crear_parrafo_estandarizado(doc, texto, negrita=False, tamano=11, espaciado_antes=0, sangria_izq=0, justificado=False, estilo=None):
        p = doc.add_paragraph()
        if estilo:
            p.style = doc.styles[estilo]
        p.paragraph_format.space_before = Pt(espaciado_antes)
        p.paragraph_format.left_indent = Inches(sangria_izq)
        if justificado:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(texto)
        run.bold = negrita
        run.font.size = Pt(tamano)
        run.font.name = 'Arial'

    crear_parrafo_estandarizado(doc, "Evaluación sumativa", negrita=True, tamano=12, espaciado_antes=18, estilo='Heading 2')

    tabla = doc.add_table(rows=7, cols=5)
    tabla.style = 'Table Grid'

    from docx.oxml import parse_xml
    
    tbl = tabla._tbl
    tblBorders = parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tbl.tblPr.append(tblBorders)

    celda_fusion = tabla.cell(0, 0)
    for col in range(1, 5):
        celda_fusion.merge(tabla.cell(0, col))
    
    celda_fusion.text = "EVALUACIÓN SUMATIVA"
    for paragraph in celda_fusion.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Arial'

    encabezados_fila2 = ["UNIDAD DE APRENDIZAJE", "EVALUACIÓN FORMATIVA", "EVALUACIÓN SUMATIVA (%)", "PESO (%)", "EVALUACIÓN"]
    for col, encabezado in enumerate(encabezados_fila2):
        celda = tabla.cell(1, col)
        celda.text = encabezado
        for paragraph in celda.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Arial'

    datos_tabla = [
        ["I", "Cuestionario", "15", "20%", "T1"],
        ["II", "Cuestionario", "15", "35%", "T2"],
        ["III", "Cuestionario", "15", "", "T3"],
        ["IV", "Cuestionario", "15", "45%", "TF"],
        ["", "", "PROMEDIO FINAL", "100%", "PF"]
    ]

    for fila_idx, datos_fila in enumerate(datos_tabla):
        for col_idx, dato in enumerate(datos_fila):
            celda = tabla.cell(fila_idx + 2, col_idx)
            celda.text = dato
            for paragraph in celda.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    if fila_idx == 4: 
                        run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'

    for col_idx in range(5):
        for fila in tabla.rows:
            celda = fila.cells[col_idx]
            if col_idx == 0:  
                celda.width = Inches(1.2)
            elif col_idx == 1:  
                celda.width = Inches(2.0)
            elif col_idx == 2: 
                celda.width = Inches(1.5)
            elif col_idx == 3:  
                celda.width = Inches(1.0)
            else:  
                celda.width = Inches(1.3)
# ------------------------------------------------------------------------------------------------------
    doc.add_paragraph("")
    crear_parrafo_estandarizado(doc, "FÓRMULA PARA LA OBTENCIÓN DE LA NOTA FINAL:", negrita=True, tamano=11, espaciado_antes=6)
    crear_parrafo_estandarizado(doc, "NF = (TA1 * 0.25) + (EC1 * 0.20) + (TA2 * 0.25) + (EC2 * 0.20) + (PA * 0.10)", tamano=10)

    parrafo = doc.add_paragraph("VII. REFERENCIAS", style='Heading 1')
    parrafo.runs[0].bold = True
    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    parrafo.paragraph_format.space_before = Pt(12)
    parrafo.runs[0].font.size = Pt(12)
    parrafo.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    referencias_cargadas = []
    try:
        with open("historial.json", "r", encoding="utf-8") as f:
            historial = json.load(f)
            # Buscar en registros_completados - Solo los últimos 5 registros para optimizar
            if "registros_completados" in historial and historial["registros_completados"]:
                registros_recientes = historial["registros_completados"][-5:]  # Solo últimos 5
                for registro in reversed(registros_recientes):
                    if "referencias" in registro:
                        referencias_data = registro["referencias"]
                        
                        # Intentar diferentes estructuras de referencias
                        # Estructura nueva: libros, articulos, web
                        if isinstance(referencias_data, dict):
                            todas_referencias = []
                            
                            # Agregar libros (máximo 10 para optimizar)
                            if "libros" in referencias_data and isinstance(referencias_data["libros"], list):
                                for libro in referencias_data["libros"][:10]:  # Límite de 10
                                    if isinstance(libro, dict):
                                        # Si ya tiene referencia formateada, usarla
                                        if "referencia_formateada" in libro:
                                            todas_referencias.append(libro)
                                        # Si no, formatear desde los campos individuales
                                        elif "titulo" in libro:
                                            try:
                                                ref_formateada = formatear_referencia_libro(libro)
                                                todas_referencias.append({"referencia_formateada": ref_formateada})
                                            except:
                                                todas_referencias.append({"referencia_formateada": "Referencia pendiente"})
                            
                            # Agregar artículos (máximo 10 para optimizar)
                            if "articulos" in referencias_data and isinstance(referencias_data["articulos"], list):
                                for articulo in referencias_data["articulos"][:10]:  # Límite de 10
                                    if isinstance(articulo, dict):
                                        # Si ya tiene referencia formateada, usarla
                                        if "referencia_formateada" in articulo:
                                            todas_referencias.append(articulo)
                                        # Si no, formatear desde los campos individuales
                                        elif "titulo" in articulo:
                                            try:
                                                ref_formateada = formatear_referencia_articulo(articulo)
                                                todas_referencias.append({"referencia_formateada": ref_formateada})
                                            except:
                                                todas_referencias.append({"referencia_formateada": "Referencia pendiente"})
                            
                            # Agregar referencias web (máximo 10 para optimizar)
                            if "web" in referencias_data and isinstance(referencias_data["web"], list):
                                for web in referencias_data["web"][:10]:  # Límite de 10
                                    if isinstance(web, dict):
                                        # Si ya tiene referencia formateada, usarla
                                        if "referencia_formateada" in web:
                                            todas_referencias.append(web)
                                        # Si no, formatear desde los campos individuales
                                        elif "titulo" in web or "url" in web:
                                            try:
                                                ref_formateada = formatear_referencia_web(web)
                                                todas_referencias.append({"referencia_formateada": ref_formateada})
                                            except:
                                                todas_referencias.append({"referencia_formateada": "Referencia pendiente"})
                            
                            if todas_referencias:
                                referencias_cargadas = todas_referencias[:10]  # Máximo 10 referencias
                                break
                        
                        # Estructura antigua: enlaces_referencia
                        elif "enlaces_referencia" in referencias_data:
                            referencias_cargadas = referencias_data["enlaces_referencia"][:10]  # Máximo 10
                            break
            
            # Si no encontramos referencias, usar valores por defecto
            if not referencias_cargadas:
                referencias_cargadas = [
                    {"referencia_formateada": "Referencia pendiente"},
                    {"referencia_formateada": "Referencia pendiente"},
                    {"referencia_formateada": "Referencia pendiente"},
                    {"referencia_formateada": "Referencia pendiente"},
                    {"referencia_formateada": "Referencia pendiente"}
                ]
                
    except (FileNotFoundError, json.JSONDecodeError, Exception) as e:
        print(f"⚠ Error al cargar referencias (usando predeterminadas): {str(e)[:100]}...")
        referencias_cargadas = [
            {"referencia_formateada": "Referencia pendiente"},
            {"referencia_formateada": "Referencia pendiente"},
            {"referencia_formateada": "Referencia pendiente"},
            {"referencia_formateada": "Referencia pendiente"},
            {"referencia_formateada": "Referencia pendiente"}
        ]

    for referencia in referencias_cargadas:
        try:
            if isinstance(referencia, dict) and "referencia_formateada" in referencia:
                texto_referencia = referencia["referencia_formateada"]
            else:
                texto_referencia = str(referencia)
            
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5) 
            
            run_bullet = p.add_run("• ")
            run_bullet.font.name = 'Times New Roman'
            run_bullet.font.size = Pt(11)
            run_bullet.italic = True
            
            run = p.add_run(texto_referencia)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.italic = True
            
        except Exception as e:
            print(f"Error al formatear la referencia: {e}")

    parrafo = doc.add_paragraph("VIII. NORMAS DE CONVIVENCIA", style='Heading 1')
    parrafo.runs[0].bold = True
    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    parrafo.paragraph_format.space_before = Pt(12)
    parrafo.runs[0].font.size = Pt(12)
    parrafo.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    normas = [
        "Respeto.",
        "Asistencia.",
        "Puntualidad.",
        "Presentación oportuna de los entregables."
    ]

    for norma in normas:
        p = doc.add_paragraph("", style='List Number')
        p.add_run(norma)
        p.paragraph_format.left_indent = Cm(1.2)

    crear_parrafo_estandarizado(doc, "Bellavista, marzo de 2025.", tamano=10, espaciado_antes=12)

    return doc


def generar_documento_word(datos_completos=None, ruta_salida=None):
    """
    Genera un documento Word con encabezado profesional para el sílabo
    
    Args:
        datos_completos (dict, optional): Diccionario con todos los datos del sílabo.
                                        Si es None, genera un documento con datos por defecto.
        ruta_salida (str, optional): Ruta donde guardar el documento. Si es None, 
                                   se genera una ruta automática.
        
    Returns:
        Document: Objeto Document de python-docx con el contenido del sílabo
    """
    if not datos_completos:
        datos_general = {
            'SLB-COD': '',
            'SLB-VER': '',
            'SLB-DATE': datetime.now().strftime('%d/%m/%Y'),
            'SLB-PROG': '',
            'SLB-ASSIG': '',
            'SLB-SEM': '',
            'SLB-DOC': ''
        }
    else:
        # Extraer datos generales
        datos_general_raw = datos_completos.get('general', {})
        
        datos_mapeados = {}
        
        for key, value in datos_general_raw.items():
            datos_mapeados[key] = value
        
        mapeo = {
            'codigo': 'SLB-COD',
            'version': 'SLB-VER', 
            'fecha': 'SLB-DATE',
            'maestria': 'SLB-PROG',
            'asignatura': 'SLB-ASSIG',
            'semestre': 'SLB-SEM',
            'docente': 'SLB-DOC',
            'sesiones': 'SESIONES',
            'semanas': 'SEMANAS',
            'correo': 'SLB-CORREO', 
            'modalidad': 'SLB-MODALIDAD',
            'horas_teoria': 'SLB-HORAT',
            'horas_practica': 'SLB-HORAP',
            'creditos': 'CRED',
            'codigo_programa': 'SLB-CODPROG',
            'caracter': 'SLB-CARACTER',
            'proposito': 'SLB-PROPOSITO'
        }
        
        for clave_flask, clave_slb in mapeo.items():
            if clave_flask in datos_general_raw and clave_slb not in datos_mapeados:
                datos_mapeados[clave_slb] = datos_general_raw[clave_flask]
        
        campos_requeridos = {
            'SLB-COD': '',
            'SLB-VER': '',
            'SLB-DATE': datetime.now().strftime('%d/%m/%Y'),
            'SLB-PROG': '',
            'SLB-ASSIG': '',
            'SLB-SEM': '',
            'SLB-DOC': '',
            'SESIONES': '16',
            'SEMANAS': '4',
            'SLB-CORREO': '',
            'SLB-MODALIDAD': 'Virtual',
            'SLB-HORAT': '0',
            'SLB-HORAP': '0',
            'CRED': '4',
            'SLB-CODPROG': '',
            'SLB-CARACTER': 'Obligatorio',
            'SLB-PROPOSITO': ''
        }
        
        for campo, valor_defecto in campos_requeridos.items():
            if campo not in datos_mapeados:
                datos_mapeados[campo] = valor_defecto
        
        # Procesar datos de unidades para crear temas_por_unidad e indicadores_por_unidad
        datos_unidades = datos_completos.get('unidades', {})
        temas_por_unidad = {}
        indicadores_por_unidad = {}
        instrumentos_por_unidad = {}
        
        # Primero intentar obtener desde datos de unidades
        if isinstance(datos_unidades, dict):
            unidades_detalle = datos_unidades.get('unidades_detalle', [])
            
            for i, unidad in enumerate(unidades_detalle, 1):
                if isinstance(unidad, dict):
                    nombre_unidad = unidad.get('nombre', f'Unidad {i}')
                    logro_unidad = unidad.get('logro', f'Desarrolla competencias específicas de la unidad {i}')
                    instrumentos_unidad = unidad.get('instrumentos', ['Lista de cotejo', 'Rúbrica de evaluación'])
                    
                    # Formatear el nombre de la unidad como clave
                    clave_unidad = f"Unidad {['I', 'II', 'III', 'IV'][i-1] if i <= 4 else str(i)}"
                    temas_por_unidad[clave_unidad] = nombre_unidad
                    indicadores_por_unidad[clave_unidad] = logro_unidad
                    instrumentos_por_unidad[clave_unidad] = instrumentos_unidad
        
        # Si no tenemos suficientes temas, intentar obtener desde datos de sesiones
        if len(temas_por_unidad) < 4:
            datos_sesiones = datos_completos.get('sesiones', {})
            if isinstance(datos_sesiones, dict) and 'unidades_sesiones' in datos_sesiones:
                for unidad_sesion in datos_sesiones['unidades_sesiones']:
                    if isinstance(unidad_sesion, dict):
                        numero_unidad = unidad_sesion.get('unidad_numero', 0)
                        nombre_unidad = unidad_sesion.get('unidad_nombre', f'Unidad {numero_unidad}')
                        if 1 <= numero_unidad <= 4:
                            clave_unidad = f"Unidad {['I', 'II', 'III', 'IV'][numero_unidad-1]}"
                            # Solo agregar si no existe ya
                            if clave_unidad not in temas_por_unidad:
                                temas_por_unidad[clave_unidad] = nombre_unidad
                                indicadores_por_unidad[clave_unidad] = f'Desarrolla competencias específicas relacionadas con {nombre_unidad.lower()}'
                                instrumentos_por_unidad[clave_unidad] = ['Lista de cotejo', 'Rúbrica de evaluación']
        
        # Agregar los datos procesados a los datos mapeados
        if temas_por_unidad:
            datos_mapeados['temas_por_unidad'] = temas_por_unidad
        if indicadores_por_unidad:
            datos_mapeados['indicadores_por_unidad'] = indicadores_por_unidad
        if instrumentos_por_unidad:
            datos_mapeados['instrumentos_por_unidad'] = instrumentos_por_unidad
        
        # Si no hay indicadores, crear valores por defecto profesionales
        if not indicadores_por_unidad:
            datos_mapeados['indicadores_por_unidad'] = {
                'Unidad I': 'Demuestra dominio de los conceptos fundamentales de la asignatura',
                'Unidad II': 'Aplica conocimientos teóricos en situaciones prácticas',
                'Unidad III': 'Integra conocimientos mediante proyectos aplicados',
                'Unidad IV': 'Evalúa y sintetiza el aprendizaje obtenido'
            }
        
        # Si no hay instrumentos, crear valores por defecto profesionales
        if not instrumentos_por_unidad:
            datos_mapeados['instrumentos_por_unidad'] = {
                'Unidad I': ['Lista de cotejo', 'Prueba de conocimientos'],
                'Unidad II': ['Rúbrica de evaluación', 'Casos prácticos'],
                'Unidad III': ['Proyecto aplicado', 'Portafolio de evidencias'],
                'Unidad IV': ['Examen integrador', 'Presentación final']
            }
        
        # Procesar datos de competencias para crear competencias_especificas
        datos_competencias = datos_completos.get('competencias', {})
        competencias_especificas = []
        
        if isinstance(datos_competencias, dict) and 'unidades_competencias' in datos_competencias:
            for unidad_comp in datos_competencias['unidades_competencias']:
                if isinstance(unidad_comp, dict) and 'competencias' in unidad_comp:
                    for comp in unidad_comp['competencias']:
                        if isinstance(comp, dict):
                            codigo = comp.get('codigo', '')
                            titulo = comp.get('titulo', '')
                            descripcion = comp.get('descripcion', '')
                            competencia_completa = f"{titulo}: {descripcion}"
                            competencias_especificas.append((codigo, competencia_completa))
        
        # Procesar datos de productos para crear productos_actividades
        datos_productos = datos_completos.get('productos', {})
        productos_actividades = []
        
        if isinstance(datos_productos, dict) and 'unidades_productos' in datos_productos:
            for unidad_prod in datos_productos['unidades_productos']:
                if isinstance(unidad_prod, dict) and 'productos' in unidad_prod:
                    for prod in unidad_prod['productos']:
                        if isinstance(prod, dict):
                            codigo = prod.get('codigo', '')
                            titulo = prod.get('titulo', '')
                            descripcion = prod.get('descripcion', '')
                            producto_completo = f"{titulo}: {descripcion}"
                            productos_actividades.append((codigo, producto_completo))
        
        # Agregar otros datos necesarios
        datos_mapeados.update({
            'competencias_especificas': competencias_especificas,
            'productos_actividades': productos_actividades,
            'cronograma': datos_completos.get('cronograma', {}),
            'referencias': datos_completos.get('referencias', {}),
            'sesiones': datos_completos.get('sesiones', {})
        })
                
        datos_general = datos_mapeados
    
    doc = crear_encabezado_profesional(datos_general)
    
    if ruta_salida is None:
        ruta_salida = os.path.join(os.path.dirname(__file__), "formato_silabo.docx")

    doc.save(ruta_salida)
    print(f"Documento guardado en: {ruta_salida}")
    
    return doc


def generar_nombre_archivo(datos_general=None):
    """
    Genera un nombre de archivo apropiado para el documento Word
    
    Args:
        datos_general (dict, optional): Diccionario con los datos del formulario general.
                                      Si es None, genera un nombre genérico.
        
    Returns:
        str: Nombre del archivo con formato: Silabo_[Asignatura]_[Fecha].docx
    """
    if datos_general and datos_general.get('asignatura'):
        asignatura = datos_general.get('asignatura', 'Silabo')
        asignatura_clean = "".join(c for c in asignatura if c.isalnum() or c in (' ', '-', '_')).rstrip()
    else:
        asignatura_clean = 'Plantilla'
    
    fecha_actual = datetime.now().strftime('%Y-%m-%d')
    return f"Silabo_{asignatura_clean}_{fecha_actual}.docx"


def obtener_fechas_unidades(cronograma_generado):
    """
    Obtiene las fechas de inicio y término de cada unidad con formato legible
    
    Args:
        cronograma_generado (dict): Diccionario con el cronograma de unidades y sesiones
        
    Returns:
        dict: Diccionario con fechas formateadas para cada unidad
    """
    print("DEBUG: Claves en cronograma_generado:", list(cronograma_generado.keys()))
    print("DEBUG: Contenido del cronograma:", cronograma_generado)
    fechas_unidades = {}
    
    for unidad, sesiones in cronograma_generado.items():
        if sesiones:
            fecha_inicio = sesiones[0]["fecha"]
            fecha_termino = sesiones[-1]["fecha"]
            
            inicio_dt = datetime.strptime(fecha_inicio, "%Y-%m-%d")
            termino_dt = datetime.strptime(fecha_termino, "%Y-%m-%d")
            
            meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio', 
                    'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
            
            fecha_inicio_fmt = f"{inicio_dt.day:02d} de {meses[inicio_dt.month-1]} de {inicio_dt.year}"
            fecha_termino_fmt = f"{termino_dt.day:02d} de {meses[termino_dt.month-1]} de {termino_dt.year}"
            
            fechas_unidades[unidad] = {
                "inicio": fecha_inicio_fmt,
                "termino": fecha_termino_fmt
            }
    
    return fechas_unidades


def obtener_fechas_sesiones(cronograma_generado):
    """
    Obtiene las fechas de todas las sesiones por unidad con formato dd/mm/yyyy
    
    Args:
        cronograma_generado (dict): Diccionario con el cronograma de unidades y sesiones
        
    Returns:
        dict: Diccionario con fechas formateadas para cada sesión por unidad
    """
    sesiones_fechas = {}
    for unidad, sesiones in cronograma_generado.items():
        sesiones_fechas[unidad] = {}
        for i, sesion in enumerate(sesiones):
            fecha_dt = datetime.strptime(sesion["fecha"], "%Y-%m-%d")
            fecha_formateada = fecha_dt.strftime("%d/%m/%Y")
            sesiones_fechas[unidad][f"sesion_{i+1}"] = fecha_formateada
    return sesiones_fechas
